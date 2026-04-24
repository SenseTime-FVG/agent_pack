"""Parse user-uploaded docs (pdf/docx/md/txt) into a normalized JSON payload.

Usage:
    python parse_user_docs.py --files path1 [path2 ...]

Output (stdout):
    {"documents": [...], "errors": [...]}
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from pypdf.errors import PdfReadError
from docx.opc.exceptions import PackageNotFoundError

MAX_CHARS = 20000


def _truncate(text: str) -> str:
    if len(text) <= MAX_CHARS:
        return text
    return text[:MAX_CHARS] + "\n[TRUNCATED]"


def parse_md_or_txt(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return {"path": str(path), "type": path.suffix.lstrip(".").lower(), "text": _truncate(raw)}


def parse_pdf(path: Path) -> dict:
    from pypdf import PdfReader
    try:
        reader = PdfReader(str(path))
        pages = [p.extract_text() or "" for p in reader.pages]
    except PdfReadError as exc:
        raise ValueError(f"corrupted pdf: {exc}") from exc
    text = "\n".join(pages)
    return {"path": str(path), "type": "pdf", "text": _truncate(text), "pages": len(reader.pages)}


def parse_docx(path: Path) -> dict:
    from docx import Document
    try:
        doc = Document(str(path))
    except PackageNotFoundError as exc:
        raise ValueError(f"corrupted or invalid docx: {exc}") from exc
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(paragraphs)
    return {"path": str(path), "type": "docx", "text": _truncate(text), "paragraphs": len(paragraphs)}


def dispatch(path: Path) -> dict:
    ext = path.suffix.lower()
    if ext in (".md", ".txt"):
        return parse_md_or_txt(path)
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".docx":
        return parse_docx(path)
    raise ValueError(f"unsupported type: {ext}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--files", nargs="+", required=True)
    args = parser.parse_args(argv)
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    documents: list[dict] = []
    errors: list[dict] = []
    for raw in args.files:
        p = Path(raw).expanduser().resolve()
        if not p.exists():
            errors.append({"path": str(p), "error": "file not found"})
            continue
        if not p.is_file():
            errors.append({"path": str(p), "error": "not a regular file"})
            continue
        try:
            documents.append(dispatch(p))
        except Exception as exc:  # noqa: BLE001
            errors.append({"path": str(p), "error": f"{type(exc).__name__}: {exc}"})

    json.dump({"documents": documents, "errors": errors}, sys.stdout, ensure_ascii=False)
    return 0


if __name__ == "__main__":
    sys.exit(main())
